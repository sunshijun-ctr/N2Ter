"""Screenplay -> Word (.docx) exporter.

Layout mirrors ``PDFExporter`` for the three schema types. Uses python-docx;
import is lazy so the app boots even when the optional dependency is missing.
"""

from __future__ import annotations

from io import BytesIO
from typing import Any

from app.services.character_profile_utils import (
    collect_character_id_map_from_episodes,
    format_shot_label,
    resolve_character_ref,
    resolve_character_text,
)


class WordExporter:
    def available(self) -> bool:
        try:
            import docx  # noqa: F401
        except ImportError:
            return False
        return True

    def render_docx(self, payload: dict[str, Any]) -> bytes:
        from docx import Document
        from docx.enum.text import WD_LINE_SPACING
        from docx.shared import Cm, Pt

        schema_type = payload.get("schema_type", "screenwriter")
        title = str(payload.get("title") or "剧本")
        episodes = payload.get("episodes") or []

        doc = Document()
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(1.8)
            section.right_margin = Cm(1.8)

        normal = doc.styles["Normal"]
        normal.font.name = "Microsoft YaHei"
        normal.font.size = Pt(11)
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        normal.paragraph_format.line_spacing = 1.5

        doc.add_heading(title, level=0)
        meta = doc.add_paragraph(
            f"格式：{schema_type} · 版本：{payload.get('schema_version', '')}"
        )
        if meta.runs:
            meta.runs[0].font.size = Pt(10)

        if schema_type == "overview":
            self._render_overview(doc, episodes)
        elif schema_type == "ai_video":
            self._render_ai_video(doc, episodes, payload.get("character_arcs") or [])
        else:
            self._render_screenwriter(doc, episodes)

        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def _add_kv(doc, label: str, value: str) -> None:
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label}：")
        run_label.bold = True
        p.add_run(str(value))

    @staticmethod
    def _add_dialogue(doc, who: str, line: str, parenthetical: str | None = None) -> None:
        from docx.shared import Cm

        p = doc.add_paragraph(style="List Paragraph")
        p.paragraph_format.left_indent = Cm(1)
        text = who
        if parenthetical:
            text += f"（{parenthetical}）"
        run_who = p.add_run(text)
        run_who.bold = True
        p.add_run(f"：{line}")

    def _render_overview(self, doc, episodes: list[dict]) -> None:
        doc_content = episodes[0].get("content") if episodes else {}
        doc_content = doc_content or {}

        if doc_content.get("logline"):
            p = doc.add_paragraph(str(doc_content["logline"]))
            p.runs[0].italic = True

        for label, key in (
            ("卖点", "hook"),
            ("题材", "genre"),
            ("预估集数", "estimated_episodes"),
            ("目标受众", "target_audience"),
            ("类比作品", "market_comparable"),
            ("改编难度", "adaptation_difficulty"),
        ):
            if doc_content.get(key):
                self._add_kv(doc, label, str(doc_content[key]))

        characters = doc_content.get("main_characters") or []
        if characters:
            doc.add_heading("主要角色", level=2)
            for character in characters:
                name = str(character.get("name", ""))
                role = str(character.get("role", ""))
                one = str(character.get("one_liner", ""))
                self._add_kv(doc, name, f"{role}　{one}".strip())

        arc = doc_content.get("plot_arc") or {}
        if arc:
            doc.add_heading("情节主线", level=2)
            for label, key in (
                ("开端", "setup"),
                ("触发", "inciting_incident"),
                ("发展", "rising_action"),
                ("高潮", "climax"),
                ("结局", "resolution"),
            ):
                if arc.get(key):
                    self._add_kv(doc, label, str(arc[key]))

        overview_episodes = doc_content.get("episodes") or []
        if overview_episodes:
            doc.add_heading("分集概览", level=2)
            for ep in overview_episodes:
                num = ep.get("episode_number", "")
                ep_title = ep.get("title", "")
                doc.add_heading(f"第 {num} 集 · {ep_title}", level=3)
                if ep.get("one_line_summary"):
                    doc.add_paragraph(str(ep["one_line_summary"]))
                self._scene_table(doc, ep.get("key_scenes") or [])

    def _scene_table(self, doc, scenes: list[dict]) -> None:
        if not scenes:
            return
        from docx.enum.table import WD_TABLE_ALIGNMENT

        table = doc.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ("地点", "出场角色", "核心冲突", "结果", "权重")
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        for scene in scenes:
            chars = scene.get("characters") or []
            chars_text = (
                "、".join(str(c) for c in chars) if isinstance(chars, list) else str(chars)
            )
            row = table.add_row().cells
            row[0].text = str(scene.get("location", ""))
            row[1].text = chars_text
            row[2].text = str(scene.get("conflict", ""))
            row[3].text = str(scene.get("outcome", ""))
            row[4].text = str(scene.get("weight", ""))

    def _render_screenwriter(self, doc, episodes: list[dict]) -> None:
        for episode in episodes:
            content = episode.get("content") or {}
            num = content.get("episode_number", episode.get("episode_number", ""))
            ep_title = content.get("title") or episode.get("title") or ""
            doc.add_heading(f"第 {num} 集 · {ep_title}", level=2)
            for label, key in (
                ("本集梗概", "episode_summary"),
                ("核心冲突", "key_conflict"),
                ("情感曲线", "emotional_arc"),
            ):
                if content.get(key):
                    self._add_kv(doc, label, str(content[key]))
            for scene in content.get("scenes") or []:
                self._render_screenwriter_scene(doc, scene)

    def _render_screenwriter_scene(self, doc, scene: dict) -> None:
        slug = scene.get("slug_line") or scene.get("scene_objective") or "场景"
        p = doc.add_paragraph()
        run = p.add_run(str(slug))
        run.bold = True
        if scene.get("action_description"):
            doc.add_paragraph(str(scene["action_description"]))
        for dialogue in scene.get("dialogues") or []:
            self._add_dialogue(
                doc,
                str(dialogue.get("character", "")),
                str(dialogue.get("line", "")),
                dialogue.get("parenthetical"),
            )
        if scene.get("rewrite_notes"):
            from docx.shared import Pt

            note = doc.add_paragraph(f"改写建议：{scene['rewrite_notes']}")
            note.runs[0].font.size = Pt(9.5)

    def _render_ai_video(
        self, doc, episodes: list[dict], character_arcs: list[dict] | None = None
    ) -> None:
        id_map = collect_character_id_map_from_episodes(episodes, character_arcs)
        for episode in episodes:
            content = episode.get("content") or {}
            num = content.get("episode_number", episode.get("episode_number", ""))
            ep_title = content.get("title") or episode.get("title") or ""
            doc.add_heading(f"第 {num} 集 · {ep_title}", level=2)
            top_shots = content.get("shots")
            if isinstance(top_shots, list) and top_shots:
                for index, shot in enumerate(top_shots, start=1):
                    self._render_ai_video_shot(
                        doc, shot, index, id_map, num, scene_index=1
                    )
            else:
                for scene_index, scene in enumerate(content.get("scenes") or [], start=1):
                    scene_label = (
                        scene.get("location")
                        or scene.get("slug_line")
                        or scene.get("heading")
                    )
                    if scene_label:
                        doc.add_heading(str(scene_label), level=3)
                    for shot_index, shot in enumerate(scene.get("shots") or [], start=1):
                        self._render_ai_video_shot(
                            doc, shot, shot_index, id_map, num, scene_index=scene_index
                        )

    def _render_ai_video_shot(
        self,
        doc,
        shot: dict,
        shot_index: int,
        id_map: dict[str, str],
        episode_num: int | str | None,
        *,
        scene_index: int,
    ) -> None:
        label = format_shot_label(
            shot,
            scene_index=scene_index,
            shot_index=shot_index,
            episode_num=episode_num,
        )
        p = doc.add_paragraph()
        run = p.add_run(label)
        run.bold = True

        for field_label, key in (
            ("景别", "shot_type"),
            ("角度", "camera_angle"),
            ("运镜", "camera_movement"),
            ("时长", "duration_seconds"),
            ("主体", "subject"),
            ("动作", "subject_action"),
            ("背景", "background"),
            ("光线", "lighting"),
            ("生成提示", "generation_prompt"),
        ):
            value = shot.get(key)
            if value in (None, ""):
                continue
            raw = str(value)
            if key == "subject" or (key == "subject_action" and "char_" in raw.lower()):
                display = resolve_character_text(raw, id_map) or resolve_character_ref(raw, id_map)
            else:
                display = raw
            self._add_kv(doc, field_label, display)

        for dialogue in shot.get("dialogue") or shot.get("dialogues") or []:
            raw_who = dialogue.get("character_id") or dialogue.get("character") or ""
            who = resolve_character_ref(str(raw_who), id_map)
            line = dialogue.get("line") or dialogue.get("text") or ""
            if not line:
                continue
            self._add_dialogue(doc, who, str(line), dialogue.get("voice_tone"))

        for emotion in shot.get("character_emotion") or []:
            if not isinstance(emotion, dict):
                continue
            raw_id = emotion.get("character_id") or emotion.get("character") or ""
            who = resolve_character_ref(str(raw_id), id_map)
            emo = emotion.get("emotion") or emotion.get("facial_expression") or ""
            if not who and not emo:
                continue
            prefix = f"{who}：" if who else ""
            self._add_kv(doc, "情绪", f"{prefix}{emo}")


word_exporter = WordExporter()
